#!/usr/bin/env python3
"""Generate Dealality GTM Word docs (separate English and Spanish files)."""

from __future__ import annotations

import glob
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
LOCAL_TEMPLATE = ROOT / "reports" / "_template-warm-intro.docx"
FONT_NAME = "Poppins"
FONT_SIZE = Pt(10)
COLOR_HEADER = RGBColor(0xC0, 0x00, 0x00)
COLOR_HEADER_ALT = RGBColor(0xEE, 0x00, 0x00)

BodyLine = tuple[str, list[str] | None]


def find_drive_folder() -> Path:
    matches = glob.glob(str(Path("G:/My Drive") / "Dealality*"))
    if not matches:
        raise FileNotFoundError("Dealality Google Drive folder not found under G:/My Drive")
    return Path(matches[0])


def ensure_local_template() -> Path:
    if LOCAL_TEMPLATE.exists():
        return LOCAL_TEMPLATE
    matches = glob.glob(str(Path("G:/My Drive") / "Dealality*" / "Dealality - Warm Introduction Templates.docx"))
    if not matches:
        raise FileNotFoundError("Template docx not found")
    LOCAL_TEMPLATE.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(matches[0], LOCAL_TEMPLATE)
    return LOCAL_TEMPLATE


def _set_run_font(run, *, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    if color is not None:
        run.font.color.rgb = color


def add_run(paragraph, text: str, *, bold: bool | None = None, color: RGBColor | None = None):
    run = paragraph.add_run(text)
    _set_run_font(run, bold=bold, color=color)
    return run


def clear_cell(cell) -> None:
    cell.text = ""


def add_section_header(cell, title: str, *, alt: bool = False) -> None:
    if cell.paragraphs:
        clear_cell(cell)
    p = cell.paragraphs[0]
    color = COLOR_HEADER_ALT if alt else COLOR_HEADER
    add_run(p, title, bold=True, color=color)


def add_subject_line(cell, subject: str, *, label: str = "Subject: ") -> None:
    p = cell.add_paragraph()
    add_run(p, label, bold=True, color=COLOR_HEADER)
    add_run(p, subject, bold=True)


def add_body_paragraph(cell, text: str, *, bold_phrases: list[str] | None = None) -> None:
    p = cell.add_paragraph()
    if not bold_phrases:
        add_run(p, text)
        return
    remaining = text
    for phrase in bold_phrases:
        if phrase not in remaining:
            continue
        before, _, after = remaining.partition(phrase)
        if before:
            add_run(p, before)
        add_run(p, phrase, bold=True)
        remaining = after
    if remaining:
        add_run(p, remaining)


def add_blank_line(cell) -> None:
    cell.add_paragraph()


def remove_all_table_rows(table) -> None:
    while len(table.rows) > 0:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)


def lines(*rows: str) -> list[BodyLine]:
    return [(row, None) for row in rows]


def pair_sections(
    title_en: str,
    title_es: str,
    subject_en: str,
    subject_es: str,
    body_en: list[BodyLine],
    body_es: list[BodyLine],
    *,
    alt: bool = False,
) -> tuple[dict, dict]:
    en = {
        "header": title_en,
        "alt": alt,
        "subject": subject_en,
        "subject_label": "Subject: ",
        "body": body_en,
    }
    es = {
        "header": title_es,
        "alt": alt,
        "subject": subject_es,
        "subject_label": "Asunto: ",
        "body": body_es,
    }
    return en, es


def add_pair(
    en_sections: list[dict],
    es_sections: list[dict],
    title_en: str,
    title_es: str,
    subject_en: str,
    subject_es: str,
    body_en: list[BodyLine],
    body_es: list[BodyLine],
    *,
    alt: bool = False,
) -> None:
    en, es = pair_sections(
        title_en, title_es, subject_en, subject_es, body_en, body_es, alt=alt
    )
    en_sections.append(en)
    es_sections.append(es)


def build_warm_intro_sections() -> tuple[list[dict], list[dict]]:
    en: list[dict] = []
    es: list[dict] = []

    add_pair(en, es,
        "PILOT WAVE 1 — STANDARD FORWARD (EMAIL)",
        "PILOTO OLA 1 — INTRO ESTÁNDAR (EMAIL)",
        "Quick intro — Joan / Dealality",
        "Introducción breve — Joan / Dealality",
        [
            ("Hi {{recipient_first_name}},", None),
            ("", None),
            (
                "I wanted to connect you with {{your_name}}, who is testing a small hospitality pilot called Dealality.",
                None,
            ),
            (
                "{{your_name}} is working with a limited group of owners and advisors to help structure hotel opportunities before brand/operator conversations — assessing readiness, comparing options, and clarifying what information may still be missing. {{owner_context}}",
                None,
            ),
            (
                "This is early and intentionally small. There is no obligation, and any real opportunity would move forward only with the owner's opt-in.",
                ["owner's opt-in"],
            ),
            ("I'll let you two take it from here if useful.", None),
            ("", None),
            ("Best,", None),
            ("{{introducer_first_name}}", None),
        ],
        [
            ("Hola {{recipient_first_name}},", None),
            ("", None),
            (
                "Quería presentarte a {{your_name}}, quien está probando un piloto pequeño de hospitalidad llamado Dealality.",
                None,
            ),
            (
                "{{your_name}} está trabajando con un grupo limitado de propietarios y asesores para ayudar a estructurar oportunidades hoteleras antes de conversaciones con marcas u operadores — evaluando preparación, comparando opciones y aclarando qué información podría faltar. {{owner_context}}",
                None,
            ),
            (
                "Esto está en una etapa temprana y es intencionalmente pequeño. No hay obligación, y cualquier oportunidad real avanzaría solo con el opt-in del propietario.",
                ["opt-in del propietario"],
            ),
            ("Los dejo conectados si les parece útil.", None),
            ("", None),
            ("Un saludo,", None),
            ("{{introducer_first_name}}", None),
        ],
    )

    add_pair(en, es,
        "PILOT WAVE 1 — SHORT FORWARD (EMAIL OR TEXT)",
        "PILOTO OLA 1 — INTRO CORTA (EMAIL O TEXTO)",
        "Quick intro — Joan / Dealality",
        "Introducción breve — Joan / Dealality",
        [
            (
                "Hi {{recipient_first_name}} — connecting you with {{your_name}}, who is running a small Dealality pilot for owners/advisors structuring hotel opportunities before brand/operator conversations. {{owner_context}} No obligation; owner opt-in only. I'll step back and let you connect if relevant.",
                ["owner opt-in only"],
            ),
            ("— {{introducer_first_name}}", None),
        ],
        [
            (
                "Hola {{recipient_first_name}} — te conecto con {{your_name}}, quien está ejecutando un piloto pequeño de Dealality para propietarios/asesores que estructuran oportunidades hoteleras antes de conversaciones con marcas u operadores. {{owner_context}} Sin obligación; solo con opt-in del propietario. Me retiro para que se conecten si es relevante.",
                ["opt-in del propietario"],
            ),
            ("— {{introducer_first_name}}", None),
        ],
        alt=True,
    )

    add_pair(en, es,
        "PILOT WAVE 1 — LINKEDIN FORWARD (DM)",
        "PILOTO OLA 1 — INTRO LINKEDIN (DM)",
        "Introduction — Joan / Dealality",
        "Introducción — Joan / Dealality",
        [
            (
                "Hi {{recipient_first_name}} — I'd like to introduce you to {{your_name}}. She's testing a small Dealality pilot to help owners and advisors structure hotel opportunities before brand/operator conversations. {{owner_context}} No obligation — only if a short conversation might be useful. Happy to connect you both here or by email.",
                None,
            ),
        ],
        [
            (
                "Hola {{recipient_first_name}} — me gustaría presentarte a {{your_name}}. Ella está probando un piloto pequeño de Dealality para ayudar a propietarios y asesores a estructurar oportunidades hoteleras antes de conversaciones con marcas u operadores. {{owner_context}} Sin obligación — solo si una conversación breve pudiera ser útil. Con gusto los conecto aquí o por email.",
                None,
            ),
        ],
    )

    add_pair(en, es,
        "PILOT WAVE 1 — JOAN REPLY (AFTER DOUBLE OPT-IN)",
        "PILOTO OLA 1 — RESPUESTA DE JOAN (DESPUÉS DE DOBLE OPT-IN)",
        "Re: Introduction — Dealality pilot",
        "Re: Introducción — piloto Dealality",
        [
            ("Hi {{recipient_first_name}},", None),
            ("", None),
            (
                "Thank you, {{introducer_first_name}}, for the introduction — and {{recipient_first_name}}, nice to meet you.",
                None,
            ),
            (
                "I'm testing a small Dealality pilot with owners and advisors in CALA and adjacent markets. The goal is practical: take one real or realistic opportunity, assess readiness, compare brand/operator alignment options, and clarify what may still be missing before conversations move too far.",
                None,
            ),
            (
                "I'm keeping the first group small. If helpful, I'd welcome a short conversation to see whether this might be relevant — whether you have an active opportunity, a realistic scenario, or simply a useful perspective.",
                None,
            ),
            (
                "Would a 20-minute call sometime in the next week or two work? I'm flexible on timing.",
                None,
            ),
            ("", None),
            ("At your service,", None),
            ("Joan", None),
        ],
        [
            ("Hola {{recipient_first_name}},", None),
            ("", None),
            (
                "Gracias, {{introducer_first_name}}, por la introducción — y {{recipient_first_name}}, un gusto saludarte.",
                None,
            ),
            (
                "Estoy probando un piloto pequeño de Dealality con propietarios y asesores en CALA y mercados adyacentes. El objetivo es práctico: tomar una oportunidad real o realista, evaluar preparación, comparar opciones de alineación marca/operador y aclarar qué podría faltar antes de que las conversaciones avancen demasiado.",
                None,
            ),
            (
                "Estoy manteniendo el primer grupo pequeño. Si es útil, me encantaría una conversación breve para ver si esto podría ser relevante — ya sea que tengas una oportunidad activa, un escenario realista o simplemente una perspectiva valiosa.",
                None,
            ),
            (
                "¿Te funcionaría una llamada de 20 minutos en la próxima semana o dos? Tengo flexibilidad de horario.",
                None,
            ),
            ("", None),
            ("A su servicio,", None),
            ("Joan", None),
        ],
        alt=True,
    )

    add_pair(en, es,
        "PLACEHOLDERS",
        "MARCADORES",
        "Reference",
        "Referencia",
        lines(
            "{{introducer_first_name}} — person making the intro",
            "{{recipient_first_name}} — person being introduced",
            "{{owner_context}} — optional one line (omit if unknown)",
            "{{your_name}} — Joan",
        ),
        lines(
            "{{introducer_first_name}} — persona que hace la introducción",
            "{{recipient_first_name}} — persona introducida",
            "{{owner_context}} — línea opcional (omitir si no se conoce)",
            "{{your_name}} — Joan",
        ),
    )

    return en, es


def build_reply_playbook_sections() -> tuple[list[dict], list[dict]]:
    templates = [
        (
            "REPLY 1 — HAPPY TO CHAT / OPEN TO A CALL",
            "RESPUESTA 1 — DISPUESTO A CONVERSAR / ABIERTO A LLAMADA",
            "Re: Dealality pilot",
            "Re: piloto Dealality",
            lines(
                "Hi {{first_name}},",
                "",
                "Thank you — I appreciate you getting back to me.",
                "",
                "I'd welcome a short conversation. I'm keeping the pilot group small, so the call is simply to see whether Dealality might be useful for a real or realistic opportunity you are advising on — or whether your perspective alone would be valuable.",
                "",
                "Would any of these work for a 20-minute call?",
                "- [Option A]",
                "- [Option B]",
                "",
                "If easier, feel free to send a couple of times that work for you.",
                "",
                "At your service,",
                "Joan",
            ),
            lines(
                "Hola {{first_name}},",
                "",
                "Gracias — aprecio mucho tu respuesta.",
                "",
                "Me encantaría una conversación breve. Estoy manteniendo el grupo piloto pequeño, así que la llamada es simplemente para ver si Dealality podría ser útil para una oportunidad real o realista que estés asesorando — o si tu perspectiva por sí sola sería valiosa.",
                "",
                "¿Alguna de estas opciones te funcionaría para una llamada de 20 minutos?",
                "- [Opción A]",
                "- [Opción B]",
                "",
                "Si es más fácil, envíame un par de horarios que te funcionen.",
                "",
                "A su servicio,",
                "Joan",
            ),
        ),
        (
            "REPLY 2 — SEND MORE INFO",
            "RESPUESTA 2 — ENVIAR MÁS INFORMACIÓN",
            "Re: Dealality pilot — overview",
            "Re: piloto Dealality — resumen",
            lines(
                "Hi {{first_name}},",
                "",
                "Happy to share a bit more.",
                "",
                "Dealality is a small pilot I'm running with owners and advisors to structure hotel opportunities before brand/operator conversations — assessing readiness, comparing alignment options, and clarifying what information may still be missing.",
                "",
                "I'm attaching / sharing the [Advisor/Consultant | Owner/Developer | Lawyer/Referral] overview — one page, no obligation.",
                "",
                "If after reading it a short conversation still seems relevant, I'd welcome that. If not, no worries at all.",
                "",
                "At your service,",
                "Joan",
            ),
            lines(
                "Hola {{first_name}},",
                "",
                "Con gusto comparto un poco más.",
                "",
                "Dealality es un piloto pequeño que estoy ejecutando con propietarios y asesores para estructurar oportunidades hoteleras antes de conversaciones con marcas u operadores — evaluando preparación, comparando opciones de alineación y aclarando qué información podría faltar.",
                "",
                "Adjunto / comparto el resumen de [Asesor/Consultor | Propietario/Desarrollador | Abogado/Referencia] — una página, sin obligación.",
                "",
                "Si después de leerlo una conversación breve sigue pareciendo relevante, me encantaría. Si no, sin problema.",
                "",
                "A su servicio,",
                "Joan",
            ),
        ),
        (
            "REPLY 3 — MAY KNOW SOMEONE / REFERRAL OFFER",
            "RESPUESTA 3 — PUEDE CONOCER A ALGUIEN / OFRECE REFERENCIA",
            "Re: Dealality pilot — introduction",
            "Re: piloto Dealality — introducción",
            lines(
                "Hi {{first_name}},",
                "",
                "That would be very helpful — thank you.",
                "",
                "I'm only looking for introductions where the owner has opted in and a conversation would genuinely be welcome. No need to share confidential details in advance.",
                "",
                "If you are open to making an intro, here is a short forwardable blurb you can paste (also happy to draft a tighter version for a specific person):",
                "",
                "[paste Standard Forward from Warm Introduction Templates]",
                "",
                "And if it is easier to start with a quick call between us first, I'm happy to do that too.",
                "",
                "At your service,",
                "Joan",
            ),
            lines(
                "Hola {{first_name}},",
                "",
                "Eso sería muy útil — gracias.",
                "",
                "Solo busco introducciones donde el propietario haya dado su opt-in y una conversación sea genuinamente bienvenida. No hace falta compartir detalles confidenciales por adelantado.",
                "",
                "Si estás abierto a hacer una intro, aquí tienes un texto corto que puedes reenviar (también con gusto redacto una versión más ajustada para una persona específica):",
                "",
                "[pegar Intro Estándar de Plantillas de Introducción]",
                "",
                "Y si es más fácil empezar con una llamada breve entre nosotros, también me funciona.",
                "",
                "A su servicio,",
                "Joan",
            ),
        ),
        (
            "REPLY 4 — CONFIDENTIALITY / DATA / CLIENT CONCERNS",
            "RESPUESTA 4 — CONFIDENCIALIDAD / DATOS / PREOCUPACIONES DEL CLIENTE",
            "Re: Dealality pilot — confidentiality",
            "Re: piloto Dealality — confidencialidad",
            lines(
                "Hi {{first_name}},",
                "",
                "Completely fair question.",
                "",
                "For this pilot, anything shared moves forward only with the owner's opt-in. I am not asking you to pass along confidential client materials. A short conversation can stay high-level until everyone is comfortable.",
                "",
                "If we proceed with a real opportunity, we would agree on what is shared, who sees it, and whether the owner wants to use the platform for structured comparison — or simply use the conversation for feedback.",
                "",
                "Happy to walk through that on a brief call if useful.",
                "",
                "At your service,",
                "Joan",
            ),
            lines(
                "Hola {{first_name}},",
                "",
                "Pregunta completamente válida.",
                "",
                "Para este piloto, cualquier cosa compartida avanza solo con el opt-in del propietario. No te estoy pidiendo que compartas materiales confidenciales de clientes. Una conversación breve puede mantenerse a alto nivel hasta que todos estén cómodos.",
                "",
                "Si avanzamos con una oportunidad real, acordaríamos qué se comparte, quién lo ve y si el propietario quiere usar la plataforma para comparación estructurada — o simplemente usar la conversación para retroalimentación.",
                "",
                "Con gusto lo revisamos en una llamada breve si es útil.",
                "",
                "A su servicio,",
                "Joan",
            ),
        ),
        (
            "REPLY 5 — ADVISORY / BROKERAGE ROLE CLARIFICATION",
            "RESPUESTA 5 — ACLARACIÓN DE ROL ASESOR / INTERMEDIACIÓN",
            "Re: Dealality pilot — advisory role",
            "Re: piloto Dealality — rol de asesoría",
            lines(
                "Hi {{first_name}},",
                "",
                "Good question — I want to be clear.",
                "",
                "Dealality is not a brokerage and is not replacing your advisory role. The pilot is about helping owners and advisors structure and compare options before brand/operator conversations — not about intermediating your client relationships.",
                "",
                "If there is a real opportunity and the owner opts in, you remain in control of how you advise and whether anything moves forward.",
                "",
                "If a short call would help clarify fit, I'm happy to do that.",
                "",
                "At your service,",
                "Joan",
            ),
            lines(
                "Hola {{first_name}},",
                "",
                "Buena pregunta — quiero ser clara.",
                "",
                "Dealality no es una correduría ni reemplaza tu rol de asesoría. El piloto se trata de ayudar a propietarios y asesores a estructurar y comparar opciones antes de conversaciones con marcas u operadores — no de intermediar tus relaciones con clientes.",
                "",
                "Si hay una oportunidad real y el propietario da su opt-in, tú sigues en control de cómo asesoras y si algo avanza.",
                "",
                "Si una llamada breve ayuda a aclarar el encaje, con gusto la hacemos.",
                "",
                "A su servicio,",
                "Joan",
            ),
        ),
        (
            "REPLY 6 — NO ACTIVE DEAL RIGHT NOW",
            "RESPUESTA 6 — SIN OPORTUNIDAD ACTIVA POR AHORA",
            "Re: Dealality pilot",
            "Re: piloto Dealality",
            lines(
                "Hi {{first_name}},",
                "",
                "Understood — thank you for letting me know.",
                "",
                "If timing changes, I'd still welcome your perspective on whether this kind of readiness workflow is useful in practice. No pressure either way.",
                "",
                "I'll check back in a few months unless you'd prefer I don't — just let me know.",
                "",
                "At your service,",
                "Joan",
            ),
            lines(
                "Hola {{first_name}},",
                "",
                "Entendido — gracias por avisarme.",
                "",
                "Si cambia el timing, igual me encantaría tu perspectiva sobre si este tipo de flujo de preparación es útil en la práctica. Sin presión en ningún sentido.",
                "",
                "Volveré a contactarte en unos meses a menos que prefieras que no — solo avísame.",
                "",
                "A su servicio,",
                "Joan",
            ),
        ),
        (
            "REPLY 7 — NOT RELEVANT / WRONG FIT",
            "RESPUESTA 7 — NO RELEVANTE / NO ES EL ENCAJE",
            "Re: Dealality pilot",
            "Re: piloto Dealality",
            lines(
                "Hi {{first_name}},",
                "",
                "Thank you for the quick reply — I appreciate it.",
                "",
                "I'll close the loop on my side. If I misread fit, apologies for the noise.",
                "",
                "If you ever come across an owner or advisor who might find this useful, a light introduction is always welcome — but absolutely no obligation.",
                "",
                "At your service,",
                "Joan",
            ),
            lines(
                "Hola {{first_name}},",
                "",
                "Gracias por la respuesta rápida — la aprecio.",
                "",
                "Cierro el ciclo de mi lado. Si malinterpreté el encaje, disculpas por el ruido.",
                "",
                "Si en algún momento encuentras un propietario o asesor a quien le pueda servir, una introducción ligera siempre es bienvenida — pero sin ninguna obligación.",
                "",
                "A su servicio,",
                "Joan",
            ),
        ),
        (
            "REPLY 8 — NO RESPONSE (FOLLOW-UP)",
            "RESPUESTA 8 — SIN RESPUESTA (SEGUIMIENTO)",
            "Re: Dealality pilot",
            "Re: piloto Dealality",
            lines(
                "Hi {{first_name}} — just wanted to follow up on my note about the Dealality pilot. No rush at all — I'm keeping the first group small and would still really value your perspective if relevant.",
                "",
                "At your service,",
                "Joan",
            ),
            lines(
                "Hola {{first_name}} — solo quería dar seguimiento a mi nota sobre el piloto Dealality. Sin prisa — estoy manteniendo el primer grupo pequeño y seguiría valorando mucho tu perspectiva si es relevante.",
                "",
                "A su servicio,",
                "Joan",
            ),
        ),
    ]

    en: list[dict] = []
    es: list[dict] = []
    for i, (title_en, title_es, subj_en, subj_es, body_en, body_es) in enumerate(templates):
        add_pair(en, es, title_en, title_es, subj_en, subj_es, body_en, body_es, alt=i % 2 == 1)
    return en, es


def build_acceptance_sections() -> tuple[list[dict], list[dict]]:
    en: list[dict] = []
    es: list[dict] = []

    add_pair(en, es,
        "PILOT ACCEPTANCE — THREE OUTCOMES",
        "ACEPTACIÓN DEL PILOTO — TRES RESULTADOS",
        "Internal reference — Pilot Wave 1",
        "Referencia interna — Piloto Ola 1",
        lines(
            "Real pilot opportunity — owner-opt-in deal or realistic scenario worth structuring in Dealality. Platform invite after access hygiene + QA.",
            "Feedback / referral only — useful perspective or future intro; no platform invite.",
            "Not a fit / defer — wrong segment, timing, or geography for Wave 1.",
        ),
        lines(
            "Oportunidad piloto real — acuerdo con opt-in del propietario o escenario realista que vale la pena estructurar en Dealality. Invitación a plataforma después de higiene de acceso + QA.",
            "Solo retroalimentación / referencia — perspectiva útil o intro futura; sin invitación a plataforma.",
            "No encaja / posponer — segmento, timing o geografía incorrectos para Ola 1.",
        ),
    )

    add_pair(en, es,
        "REAL PILOT OPPORTUNITY — REQUIRED (ALL 5)",
        "OPORTUNIDAD PILOTO REAL — REQUISITOS (LOS 5)",
        "Acceptance checklist",
        "Lista de aceptación",
        lines(
            "1. Owner opt-in — owner or authorized decision-maker agreed to explore Dealality for a specific opportunity or realistic scenario.",
            "2. Defined subject — identifiable hotel project: location, asset type/scale, and stage.",
            "3. Pilot-appropriate scope — one opportunity, not portfolio-wide or vague exploration.",
            "4. Confidentiality comfort — participant understands what is shared and who sees it.",
            "5. Joan capacity — supportable within Wave 1 (typically ≤3 active real opportunities).",
        ),
        lines(
            "1. Opt-in del propietario — el propietario o tomador de decisiones autorizado acordó explorar Dealality para una oportunidad específica o escenario realista.",
            "2. Sujeto definido — proyecto hotelero identificable: ubicación, tipo/escala del activo y etapa.",
            "3. Alcance apropiado para piloto — una oportunidad, no cartera completa ni exploración vaga.",
            "4. Comodidad con confidencialidad — el participante entiende qué se comparte y quién lo ve.",
            "5. Capacidad de Joan — manejable dentro de Ola 1 (típicamente ≤3 oportunidades reales activas).",
        ),
        alt=True,
    )

    add_pair(en, es,
        "REAL PILOT OPPORTUNITY — SUPPORTING SIGNALS (≥2)",
        "OPORTUNIDAD PILOTO REAL — SEÑALES DE APOYO (≥2)",
        "Supporting signals",
        "Señales de apoyo",
        lines(
            "Timing: decision, LOI, RFP, or brand/operator conversation within ~12 months.",
            "Objective: reflag, new build, operator search, brand comparison, or restructuring.",
            "Information depth: willing to share enough for readiness assessment.",
            "Advisor alignment: lawyer/advisor supports structured comparison, or owner is self-advised.",
            "CALA / pilot geography fit: aligns with Wave 1 focus unless explicitly prioritized.",
        ),
        lines(
            "Timing: decisión, LOI, RFP o conversación con marca/operador dentro de ~12 meses.",
            "Objetivo: reflag, desarrollo nuevo, búsqueda de operador, comparación de marcas o reestructuración.",
            "Profundidad de información: dispuesto a compartir suficiente para evaluación de preparación.",
            "Alineación del asesor: abogado/asesor apoya comparación estructurada, o propietario se asesora solo.",
            "Encaje geográfico CALA / piloto: alineado con enfoque Ola 1 salvo priorización explícita.",
        ),
    )

    add_pair(en, es,
        "DISQUALIFIERS (ANY ONE → NOT REAL PILOT YET)",
        "DESCALIFICADORES (CUALQUIERA → AÚN NO ES PILOTO REAL)",
        "Disqualifiers",
        "Descalificadores",
        lines(
            "No owner opt-in; advisor-only curiosity with no client engagement path.",
            "Request to use Dealality as broker/intermediary for confidential pipelines.",
            "Brand or operator asking for owner lead flow.",
            "Unwilling to share even high-level opportunity parameters.",
            "Purely theoretical with no path to real scenario in 12 months.",
        ),
        lines(
            "Sin opt-in del propietario; curiosidad solo del asesor sin ruta de engagement con cliente.",
            "Solicitud de usar Dealality como corredor/intermediario de pipelines confidenciales.",
            "Marca u operador pidiendo flujo de leads de propietarios.",
            "No dispuesto a compartir ni parámetros de alto nivel de la oportunidad.",
            "Puramente teórico sin ruta a escenario real en 12 meses.",
        ),
        alt=True,
    )

    add_pair(en, es,
        "ON-CALL CONFIRMATION QUESTIONS",
        "PREGUNTAS DE CONFIRMACIÓN EN LLAMADA",
        "Use on every pilot call",
        "Usar en cada llamada piloto",
        lines(
            '1. "Is there a specific opportunity or realistic scenario you would want to structure — or is this mainly perspective / referral?"',
            '2. "If we proceed, anything shared moves forward only with the owner\'s opt-in — does that work?"',
            '3. "What would make this a useful use of 30 minutes for you?"',
        ),
        lines(
            '1. "¿Hay una oportunidad específica o escenario realista que quisieras estructurar — o esto es principalmente perspectiva / referencia?"',
            '2. "Si avanzamos, cualquier cosa compartida sigue solo con el opt-in del propietario — ¿te funciona?"',
            '3. "¿Qué haría que estos 30 minutos sean útiles para ti?"',
        ),
    )

    add_pair(en, es,
        "WAVE 1 SUCCESS METRICS",
        "MÉTRICAS DE ÉXITO OLA 1",
        "Founder targets",
        "Metas del fundador",
        lines(
            "Real pilot opportunities accepted: 1–3",
            "Platform invites sent: only accepted real opportunities",
            "Feedback / referral conversations: unlimited; prioritize quality notes",
        ),
        lines(
            "Oportunidades piloto reales aceptadas: 1–3",
            "Invitaciones a plataforma enviadas: solo oportunidades reales aceptadas",
            "Conversaciones de retroalimentación / referencia: ilimitadas; priorizar notas de calidad",
        ),
        alt=True,
    )

    return en, es


def build_pilot_call_script_sections() -> tuple[list[dict], list[dict]]:
    en: list[dict] = []
    es: list[dict] = []

    add_pair(en, es,
        "BEFORE THE CALL — PREP",
        "ANTES DE LA LLAMADA — PREPARACIÓN",
        "Internal checklist",
        "Lista interna",
        lines(
            "Confirm segment: Lawyer/Advisor, Consultant/Broker, or Owner/Developer.",
            "Review Pilot Target List row: segment, reply notes, warm intro path.",
            "Have correct pilot overview one-pager ready if asked.",
            "Do not share platform access unless acceptance criteria are met.",
        ),
        lines(
            "Confirmar segmento: Abogado/Asesor, Consultor/Broker o Propietario/Desarrollador.",
            "Revisar fila en Pilot Target List: segmento, notas de respuesta, ruta de intro.",
            "Tener listo el resumen piloto correcto si lo piden.",
            "No compartir acceso a plataforma salvo que se cumplan criterios de aceptación.",
        ),
    )

    add_pair(en, es,
        "OPENING — ALL SEGMENTS (20–30 MIN)",
        "APERTURA — TODOS LOS SEGMENTOS (20–30 MIN)",
        "Call framing",
        "Marco de la llamada",
        lines(
            "Thanks for making time — I really appreciate it.",
            "",
            "I'm testing a small Dealality pilot with a limited group of owners and advisors. The goal today is simple: understand whether this might be useful for you — whether that's an active opportunity, a realistic scenario, feedback on the workflow, or a future introduction.",
            "",
            "This is early-stage and low-pressure. Nothing moves forward without the owner's opt-in, and we can keep things high-level until everyone is comfortable.",
            "",
            "Does that framing work for you?",
        ),
        lines(
            "Gracias por tu tiempo — lo aprecio mucho.",
            "",
            "Estoy probando un piloto pequeño de Dealality con un grupo limitado de propietarios y asesores. El objetivo hoy es simple: entender si esto podría ser útil para ti — ya sea una oportunidad activa, un escenario realista, retroalimentación sobre el flujo o una introducción futura.",
            "",
            "Esto está en etapa temprana y sin presión. Nada avanza sin el opt-in del propietario, y podemos mantener las cosas a alto nivel hasta que todos estén cómodos.",
            "",
            "¿Te funciona ese marco?",
        ),
        alt=True,
    )

    add_pair(en, es,
        "LAWYER / ADVISOR BRANCH",
        "RAMA ABOGADO / ASESOR",
        "Discovery focus",
        "Enfoque de descubrimiento",
        lines(
            "You often see deals before they're fully structured for brand or operator conversations. Dealality helps owners and advisors assess readiness, compare alignment options, and see what may still be missing — before outreach goes too far.",
            "",
            "1. Are you advising an owner on a hotel opportunity today — or mainly exploring whether this structure would be useful in your practice?",
            "2. If there is a client situation, is the owner open to a structured comparison workflow (with their opt-in)?",
            "3. What usually goes wrong when owners approach brands/operators too early?",
            "4. Would a short pilot on one real or realistic scenario be useful — or is feedback on the concept enough for now?",
            "",
            "If referral only: offer warm intro blurb; do not push platform access.",
        ),
        lines(
            "A menudo ves operaciones antes de que estén estructuradas para conversaciones con marcas u operadores. Dealality ayuda a propietarios y asesores a evaluar preparación, comparar opciones de alineación y ver qué puede faltar — antes de que el outreach avance demasiado.",
            "",
            "1. ¿Estás asesorando a un propietario en una oportunidad hotelera hoy — o principalmente explorando si esta estructura sería útil en tu práctica?",
            "2. Si hay un caso de cliente, ¿el propietario está abierto a un flujo de comparación estructurada (con su opt-in)?",
            "3. ¿Qué suele salir mal cuando los propietarios se acercan a marcas/operadores demasiado pronto?",
            "4. ¿Un piloto corto con un escenario real o realista sería útil — o basta retroalimentación sobre el concepto?",
            "",
            "Si solo referencia: ofrecer blurb de intro; no presionar acceso a plataforma.",
        ),
    )

    add_pair(en, es,
        "CONSULTANT / BROKER BRANCH",
        "RAMA CONSULTOR / BROKER",
        "Discovery focus",
        "Enfoque de descubrimiento",
        lines(
            "You sit close to owner conversations about positioning, timing, and who to approach. The pilot is about structuring and comparing options before those conversations — not replacing your role.",
            "",
            "1. Do you have an owner or project in mind where readiness or comparison would help — or is this about whether the workflow fits how you advise?",
            "2. How do you usually help owners prepare before brand/operator outreach?",
            "3. What information do owners typically lack at that stage?",
            "4. If relevant, would the owner be open to exploring this with their opt-in — or prefer feedback/referral for now?",
            "",
            "Clarify if needed: Dealality is not a brokerage; you remain in control of client relationships.",
        ),
        lines(
            "Estás cerca de conversaciones de propietarios sobre posicionamiento, timing y a quién acercarse. El piloto trata de estructurar y comparar opciones antes de esas conversaciones — no de reemplazar tu rol.",
            "",
            "1. ¿Tienes un propietario o proyecto en mente donde preparación o comparación ayudaría — o se trata de si el flujo encaja con cómo asesoras?",
            "2. ¿Cómo sueles ayudar a propietarios a prepararse antes del outreach con marcas/operadores?",
            "3. ¿Qué información suelen carecer los propietarios en esa etapa?",
            "4. Si aplica, ¿el propietario estaría abierto a explorar esto con su opt-in — o prefieren retroalimentación/referencia por ahora?",
            "",
            "Aclarar si hace falta: Dealality no es correduría; tú sigues en control de las relaciones con clientes.",
        ),
        alt=True,
    )

    add_pair(en, es,
        "OWNER / DEVELOPER BRANCH",
        "RAMA PROPIETARIO / DESARROLLADOR",
        "Discovery focus",
        "Enfoque de descubrimiento",
        lines(
            "Dealality helps owners structure one opportunity at a time — assess readiness, compare brand/operator alignment, and clarify gaps before conversations move too far. You stay in control of what is shared.",
            "",
            "1. Is there a specific project or scenario you'd want to work through — location, asset type, stage?",
            "2. What are you trying to decide in the next 6–12 months?",
            "3. What do you already know vs. what still feels unclear before talking to brands or operators?",
            "4. Who else is involved (advisor, lawyer, capital) and would they need to be part of a pilot?",
            "",
            "Early-stage: realistic hypothetical OK for feedback; real pilot needs defined subject + opt-in.",
        ),
        lines(
            "Dealality ayuda a propietarios a estructurar una oportunidad a la vez — evaluar preparación, comparar alineación marca/operador y aclarar brechas antes de que las conversaciones avancen. Tú controlas qué se comparte.",
            "",
            "1. ¿Hay un proyecto o escenario específico que quisieras trabajar — ubicación, tipo de activo, etapa?",
            "2. ¿Qué intentas decidir en los próximos 6–12 meses?",
            "3. ¿Qué ya sabes vs. qué aún se siente poco claro antes de hablar con marcas u operadores?",
            "4. ¿Quién más participa (asesor, abogado, capital) y deberían ser parte de un piloto?",
            "",
            "Etapa temprana: hipotético realista OK para retroalimentación; piloto real requiere sujeto definido + opt-in.",
        ),
    )

    add_pair(en, es,
        "CORE DISCOVERY QUESTIONS (ANY SEGMENT)",
        "PREGUNTAS CLAVE DE DESCUBRIMIENTO (CUALQUIER SEGMENTO)",
        "Use conversationally — not as a rigid checklist",
        "Usar de forma conversacional — no como lista rígida",
        lines(
            "Subject: One primary opportunity/scenario, or mainly general perspective?",
            "Timing: Decision, LOI, RFP, or brand/operator conversation on the horizon?",
            "Objective: Reflag, new build, operator search, brand comparison, restructuring?",
            "Geography: Where is the project? CALA vs other?",
            "Information: What are you comfortable sharing at a high level today?",
            "Success: What would make this call useful for you in 30 minutes?",
        ),
        lines(
            "Sujeto: ¿Una oportunidad/escenario principal, o principalmente perspectiva general?",
            "Timing: ¿Decisión, LOI, RFP o conversación con marca/operador en el horizonte?",
            "Objetivo: ¿Reflag, desarrollo nuevo, búsqueda de operador, comparación de marcas, reestructuración?",
            "Geografía: ¿Dónde está el proyecto? ¿CALA u otro?",
            "Información: ¿Qué te sentirías cómodo compartiendo a alto nivel hoy?",
            "Éxito: ¿Qué haría útil esta llamada de 30 minutos para ti?",
        ),
        alt=True,
    )

    add_pair(en, es,
        "FEEDBACK VS REAL OPPORTUNITY (ON CALL)",
        "RETROALIMENTACIÓN VS OPORTUNIDAD REAL (EN LLAMADA)",
        "Decision branch — align with acceptance criteria",
        "Rama de decisión — alinear con criterios de aceptación",
        lines(
            '1. "Is there a specific opportunity or realistic scenario you would want to structure — or is this mainly perspective / referral?"',
            '2. "If we proceed, anything shared moves forward only with the owner\'s opt-in — does that work?"',
            '3. "What would make this a useful use of our time?"',
            "",
            "Owner opt-in + defined subject → Real pilot opportunity (note gaps; schedule intake follow-up).",
            "Useful perspective, no active deal → Feedback / referral only.",
            "Wrong timing or fit → Follow-up later or not a fit.",
        ),
        lines(
            '1. "¿Hay una oportunidad específica o escenario realista que quisieras estructurar — o esto es principalmente perspectiva / referencia?"',
            '2. "Si avanzamos, cualquier cosa compartida sigue solo con el opt-in del propietario — ¿te funciona?"',
            '3. "¿Qué haría útil nuestro tiempo?"',
            "",
            "Opt-in del propietario + sujeto definido → Oportunidad piloto real (anotar brechas; agendar seguimiento de intake).",
            "Perspectiva útil, sin operación activa → Solo retroalimentación / referencia.",
            "Timing o encaje incorrecto → Seguimiento después o no encaja.",
        ),
    )

    add_pair(en, es,
        "CONFIDENTIALITY & OPT-IN (SCRIPTED)",
        "CONFIDENCIALIDAD Y OPT-IN (GUIÓN)",
        "Say before sensitive topics",
        "Decir antes de temas sensibles",
        lines(
            "Just to be clear: I'm not asking you to share confidential materials on this call unless you're comfortable.",
            "",
            "If we explore a real opportunity later, we agree on what's shared, who sees it, and whether you want structured comparison on the platform — or keep it conversational.",
            "",
            "Nothing goes to brands or operators without your control.",
        ),
        lines(
            "Para ser clara: no te pido compartir materiales confidenciales en esta llamada a menos que te sientas cómodo.",
            "",
            "Si exploramos una oportunidad real después, acordamos qué se comparte, quién lo ve y si quieres comparación estructurada en la plataforma — o mantenerlo conversacional.",
            "",
            "Nada va a marcas u operadores sin tu control.",
        ),
        alt=True,
    )

    add_pair(en, es,
        "CLOSE & NEXT STEPS — REAL PILOT CANDIDATE",
        "CIERRE Y PRÓXIMOS PASOS — CANDIDATO PILOTO REAL",
        "If acceptance criteria trending yes",
        "Si criterios de aceptación van hacia sí",
        lines(
            "This sounds like it could be a good fit for the pilot.",
            "",
            "Next step would be a short follow-up to capture the basics we discussed and see if structured comparison would help — still with your opt-in.",
            "",
            "I'll send a brief summary and we can pick a time.",
        ),
        lines(
            "Esto suena como que podría encajar bien en el piloto.",
            "",
            "El siguiente paso sería un seguimiento breve para capturar lo básico que comentamos y ver si la comparación estructurada ayudaría — siempre con tu opt-in.",
            "",
            "Enviaré un resumen breve y podemos coordinar horario.",
        ),
    )

    add_pair(en, es,
        "CLOSE & NEXT STEPS — FEEDBACK / REFERRAL ONLY",
        "CIERRE Y PRÓXIMOS PASOS — SOLO RETROALIMENTACIÓN / REFERENCIA",
        "If not a real pilot yet",
        "Si aún no es piloto real",
        lines(
            "This was really helpful — thank you.",
            "",
            "I'll follow up with the one-pager if useful, and if you think of an owner who'd welcome a conversation, I can send a short intro blurb you can forward.",
            "",
            "Thanks again for your time.",
        ),
        lines(
            "Esto fue muy útil — gracias.",
            "",
            "Haré seguimiento con el resumen de una página si sirve, y si se te ocurre un propietario que quiera conversar, puedo enviar un texto corto de intro para reenviar.",
            "",
            "Gracias de nuevo por tu tiempo.",
        ),
        alt=True,
    )

    add_pair(en, es,
        "POST-CALL — PILOT TARGET LIST",
        "POST-LLAMADA — PILOT TARGET LIST",
        "Log within 24 hours",
        "Registrar dentro de 24 horas",
        lines(
            "Reply Notes: date, call summary, outcome path, concerns raised.",
            "Pilot Fit: per Pilot Acceptance Criteria document.",
            "Outreach Status: Meeting Scheduled; Converted To Pilot only when accepted.",
            "Next Follow-Up Date: if defer or waiting on owner opt-in.",
            "Warm Intro? / Warm Intro Contact: if referral path discussed.",
        ),
        lines(
            "Reply Notes: fecha, resumen de llamada, ruta de resultado, preocupaciones.",
            "Pilot Fit: según documento de Criterios de Aceptación del Piloto.",
            "Outreach Status: Meeting Scheduled; Converted To Pilot solo cuando se acepte.",
            "Next Follow-Up Date: si se pospone o se espera opt-in del propietario.",
            "Warm Intro? / Warm Intro Contact: si se discutió ruta de referencia.",
        ),
    )

    return en, es


def build_daniel_shamah_call_package_sections() -> tuple[list[dict], list[dict]]:
    en: list[dict] = []
    es: list[dict] = []

    add_pair(en, es,
        "CALL OVERVIEW",
        "RESUMEN DE LA LLAMADA",
        "Daniel Shamah — Thu Jul 9, 2026 · 30 min",
        "Daniel Shamah — jue 9 jul 2026 · 30 min",
        lines(
            "Contact: Daniel Shamah — Partner, ECIJA Panama · Lawyer · Spanish",
            "Time: Thu Jul 9, 11:00 AM Panama / 6:00 PM Madrid",
            "Intro: Gabriel (ECIJA) · Pilot Target List recneug0VzApaESaZ",
            "Goal: Learn — not sell. Validate workflow, trust, confidentiality, referral appetite.",
            "Expected Pilot Fit: Feedback / Referral Only (unless real client scenario + owner opt-in).",
        ),
        lines(
            "Contacto: Daniel Shamah — Partner, ECIJA Panamá · Abogado · Español",
            "Hora: jue 9 jul, 11:00 AM Panamá / 6:00 PM Madrid",
            "Intro: Gabriel (ECIJA) · Pilot Target List recneug0VzApaESaZ",
            "Objetivo: Aprender — no vender. Validar flujo, confianza, confidencialidad, disposición a referir.",
            "Pilot Fit esperado: Solo retroalimentación / referencia (salvo caso real de cliente + opt-in del propietario).",
        ),
    )

    add_pair(en, es,
        "BEFORE THE CALL — INTERNAL PREP",
        "ANTES DE LA LLAMADA — PREPARACIÓN INTERNA",
        "Checklist",
        "Lista de verificación",
        lines(
            "Open demo tabs; dry-run once in Spanish.",
            "Have Lawyer / Advisor one-pager ready if asked.",
            "Have warm intro blurb ready if referral path opens.",
            "Do not grant platform login on this call.",
            "Log notes in Pilot Target List → Reply Notes after the call.",
        ),
        lines(
            "Abrir pestañas de demo; ensayo en seco una vez en español.",
            "Tener listo el resumen Abogado / Asesor si lo piden.",
            "Tener listo el blurb de intro cálida si se abre vía de referencia.",
            "No otorgar acceso/login a la plataforma en esta llamada.",
            "Registrar notas en Pilot Target List → Reply Notes después de la llamada.",
        ),
        alt=True,
    )

    add_pair(en, es,
        "30-MINUTE AGENDA",
        "AGENDA DE 30 MINUTOS",
        "Timed blocks",
        "Bloques cronometrados",
        lines(
            "0:00–2:00 — Welcome + confidentiality frame",
            "2:00–7:00 — Dealality framing (5 min)",
            "7:00–17:00 — Demo walkthrough (10 min)",
            "17:00–27:00 — Feedback questions (10 min)",
            "27:00–30:00 — Referral + next step (3 min)",
        ),
        lines(
            "0:00–2:00 — Bienvenida + marco de confidencialidad",
            "2:00–7:00 — Encuadre de Dealality (5 min)",
            "7:00–17:00 — Demostración (10 min)",
            "17:00–27:00 — Preguntas de retroalimentación (10 min)",
            "27:00–30:00 — Referencia + siguiente paso (3 min)",
        ),
    )

    add_pair(en, es,
        "OPENING + CONFIDENTIALITY (0:00–2:00)",
        "APERTURA + CONFIDENCIALIDAD (0:00–2:00)",
        "Say in Spanish (Daniel's language)",
        "Decir en español (idioma de Daniel)",
        lines(
            "Daniel, thank you for your time — and thanks to Gabriel for the connection.",
            "",
            "As I mentioned by email, this is not a sales pitch. I'm in an early validation stage with a small group of professionals who understand how these decisions are really made in hospitality.",
            "",
            "Today I'd like to briefly show you what Dealality looks like and, above all, hear your perspective: whether the workflow makes sense for owners and advisors, what legal or confidentiality concerns you'd see, and whether an introduction might make sense someday — only with the owner's opt-in.",
            "",
            "You don't need to share confidential client information on this call. Does that framing work for you?",
        ),
        lines(
            "Daniel, gracias por tu tiempo — y gracias a Gabriel por la conexión.",
            "",
            "Como comenté por email, esto no es una presentación comercial. Estoy en una etapa temprana de validación con un grupo pequeño de profesionales que entienden cómo se toman estas decisiones en hotelería.",
            "",
            "Hoy me gustaría mostrarte brevemente cómo se ve Dealality y, sobre todo, escuchar tu perspectiva: si el flujo tiene sentido para propietarios y asesores, qué preocupaciones legales o de confidencialidad verías, y si en algún momento tendría sentido una introducción — solo con el opt-in del propietario.",
            "",
            "No necesitas compartir información confidencial de clientes en esta llamada. ¿Te funciona ese marco?",
        ),
        alt=True,
    )

    add_pair(en, es,
        "5-MINUTE DEALALITY FRAMING (2:00–7:00)",
        "ENCUADRE DE DEALALITY — 5 MINUTOS (2:00–7:00)",
        "What it is / is not",
        "Qué es / qué no es",
        lines(
            "Dealality is a confidential, owner-controlled workspace for one hotel opportunity at a time.",
            "",
            "It helps organize the process before brand/operator conversations scatter: what's clear, what's missing, who should participate, what to compare, and the next step.",
            "",
            "It does not replace lawyers, advisors, or commercial relationships. Trusted advisors stay in the circle — the owner arrives better prepared.",
            "",
            "It is NOT: a brokerage, a lead pipeline for brands/operators, or a commitment to share anything without owner opt-in.",
            "",
            "Why Daniel's view matters: ECIJA sees hotel, tourism, and real estate deals in Central America before they're ready for brand/operator outreach.",
            "",
            "Transition: If you're OK with it, I'll show a 10-minute demo with a sample opportunity in Mexico — no real client data — then I'd love your feedback.",
        ),
        lines(
            "Dealality es una capa de trabajo confidencial, controlada por el propietario, para una oportunidad hotelera a la vez.",
            "",
            "Ayuda a ordenar el proceso antes de que las conversaciones con marcas u operadores se dispersen: qué ya está claro, qué falta, quién debe participar, qué comparar, y cuál es el siguiente paso.",
            "",
            "No reemplaza abogados, asesores ni la relación comercial. Ustedes siguen siendo el círculo de confianza — y el propietario llega mejor preparado.",
            "",
            "NO es: correduría, pipeline de leads para marcas/operadores, ni compromiso de compartir nada sin opt-in del propietario.",
            "",
            "Por qué importa la perspectiva de Daniel: ECIJA ve operaciones hoteleras, turismo e inmobiliario en Centroamérica antes de que estén listas para marca/operador.",
            "",
            "Transición: Si te parece, te muestro en 10 minutos un ejemplo de demostración en México — sin datos de clientes reales — y luego me gustaría tu retroalimentación.",
        ),
    )

    add_pair(en, es,
        "10-MINUTE DEMO WALKTHROUGH (7:00–17:00)",
        "DEMOSTRACIÓN — 10 MINUTOS (7:00–17:00)",
        "Screen-share · sample CALA deal only",
        "Compartir pantalla · solo oportunidad demo CALA",
        lines(
            "Demo deal: Mérida Centro Select-Service (CALA sample — not a real client).",
            "Mode: Screen-share only. No login. No Airtable edits on the call.",
            "",
            "Step 1 (~1 min) — Home / dashboard: entry point; active opportunity; readiness; next steps.",
            "Step 2 (~2 min) — Deal Setup: structured basics — location, asset, stage, objective.",
            "Step 3 (~2 min) — Brand Explorer: brand option comparison with alignment context — reference, not auto-recommendation.",
            "Step 4 (~2 min) — Operator Explorer: same logic for operators — model, geography, asset fit.",
            "Step 5 (~2 min) — Deal Room / compare: one place to view opportunity and comparisons — owner controls what is shared.",
            "Step 6 (~1 min) — Pause: This is demo data. In a real case, everything moves only with owner opt-in.",
            "",
            "Guardrail: Nothing goes to brands/operators without owner control. Trusted advisors can stay outside or inside the flow per client agreement.",
        ),
        lines(
            "Oportunidad demo: Mérida Centro Select-Service (ejemplo CALA — no es cliente real).",
            "Modo: Solo compartir pantalla. Sin login. Sin ediciones en Airtable durante la llamada.",
            "",
            "Paso 1 (~1 min) — Inicio / dashboard: punto de entrada; oportunidad activa; preparación; próximos pasos.",
            "Paso 2 (~2 min) — Deal Setup: lo básico estructurado — ubicación, activo, etapa, objetivo.",
            "Paso 3 (~2 min) — Brand Explorer: comparación de marcas con contexto de alineación — referencia, no recomendación automática.",
            "Paso 4 (~2 min) — Operator Explorer: misma lógica para operadores — modelo, geografía, encaje.",
            "Paso 5 (~2 min) — Deal Room / comparar: un lugar para ver la oportunidad y comparaciones — el propietario controla qué se comparte.",
            "Paso 6 (~1 min) — Pausa: Esto es demostración con datos de ejemplo. En un caso real, todo avanza solo con opt-in del propietario.",
            "",
            "Salvaguarda: Nada va a marcas/operadores sin control del propietario. Los asesores de confianza pueden quedar fuera o dentro del flujo según acuerden con su cliente.",
        ),
        alt=True,
    )

    add_pair(en, es,
        "10-MINUTE FEEDBACK QUESTIONS (17:00–27:00)",
        "PREGUNTAS DE RETROALIMENTACIÓN — 10 MIN (17:00–27:00)",
        "Four learning goals — listen more than talk",
        "Cuatro objetivos de aprendizaje — escuchar más que hablar",
        lines(
            "1. WORKFLOW — Does this make sense for owners/advisors?",
            "   • Does ordering the opportunity before brand/operator outreach match how your hotel clients decide?",
            "   • When in the process would a Central America owner benefit most?",
            "   • What feels useful vs. unnecessary?",
            "",
            "2. TRUST — Would lawyers/advisors trust this process?",
            "   • What would need to be true for you to be comfortable recommending a client explore this?",
            "   • What role should the trusted lawyer/advisor play?",
            "   • What would make you say no, this complicates the client relationship?",
            "",
            "3. CONFIDENTIALITY / LEGAL",
            "   • Confidentiality, privilege, or conflict concerns in the region?",
            "   • What disclaimers or safeguards would you expect before a owner shares information?",
            "   • Local practices (Panama / Central America) that should shape the design?",
            "",
            "4. REFERRAL (if demo feels credible)",
            "   • If credible for your clients, who might you introduce — owner, advisor, operator?",
            "   • Informal intro, forwardable email, or wait for a concrete case?",
            "   • Anyone in mind today — or mainly general perspective?",
            "",
            "Decision: Feedback/referral only → thank + one-pager + warm intro blurb.",
            "Possible pilot → intake follow-up with owner opt-in; no login today.",
        ),
        lines(
            "1. FLUJO — ¿Tiene sentido para propietarios/asesores?",
            "   • ¿Ordenar la oportunidad antes del outreach con marca/operador refleja cómo deciden sus clientes hoteleros?",
            "   • ¿En qué momento se beneficiaría más un propietario en Centroamérica?",
            "   • ¿Qué parte se siente útil vs. innecesaria?",
            "",
            "2. CONFIANZA — ¿Confiarían abogados/asesores en este proceso?",
            "   • ¿Qué tendría que ser cierto para recomendar que un cliente explore esto?",
            "   • ¿Qué rol debería tener el abogado/asesor de confianza?",
            "   • ¿Qué les haría decir no, esto complica la relación con el cliente?",
            "",
            "3. CONFIDENCIALIDAD / LEGAL",
            "   • ¿Preocupaciones de confidencialidad, privilegio o conflicto en la región?",
            "   • ¿Qué disclaimers o salvaguardas esperarían antes de que un propietario comparta información?",
            "   • ¿Prácticas locales (Panamá / Centroamérica) que cambiarían el diseño?",
            "",
            "4. REFERENCIA (si la demo se siente creíble)",
            "   • Si es creíble para sus clientes, ¿a quién introducir — propietario, asesor, operador?",
            "   • ¿Intro informal, email reenviable, o esperar un caso concreto?",
            "   • ¿Alguien en mente hoy — o más perspectiva general?",
            "",
            "Decisión: Solo retroalimentación/referencia → agradecer + resumen de una página + blurb de intro.",
            "Posible piloto → seguimiento de intake con opt-in del propietario; sin login hoy.",
        ),
    )

    add_pair(en, es,
        "5-MINUTE REFERRAL + NEXT STEP (27:00–30:00)",
        "REFERENCIA + SIGUIENTE PASO — 5 MIN (27:00–30:00)",
        "One clear ask at close",
        "Una sola petición clara al cerrar",
        lines(
            "If feedback-only (most likely):",
            "Daniel, this was very helpful — thank you. I can send a one-page summary you can forward if an interested owner comes up. If you ever see a client who would benefit from ordering the process before brand/operator conversations — with their opt-in — I'd welcome an introduction.",
            "",
            "If he offers a specific intro:",
            "I'll send a short forwardable paragraph; I'll follow the owner's timing. Nothing moves without opt-in.",
            "",
            "If legal concerns to address later:",
            "Very valuable. I can send a draft confidentiality framing for email review.",
            "",
            "Close: Thank you again — regards to Gabriel. I'll follow up with one clear next step.",
        ),
        lines(
            "Si solo retroalimentación (lo más probable):",
            "Daniel, esto fue muy útil — gracias de verdad. Si te parece, te envío un resumen de una página que puedes reenviar si surge un propietario interesado. Si en algún momento ves un cliente al que le serviría ordenar el proceso antes de hablar con marcas u operadores — con su opt-in — me encantaría la introducción.",
            "",
            "Si ofrece una intro específica:",
            "Te envío un párrafo corto que puedes reenviar; me adapto al timing del propietario. Nada avanza sin su opt-in.",
            "",
            "Si plantea temas legales para después:",
            "Muy valioso. ¿Te parece si te envío un borrador del marco de confidencialidad para revisar por email?",
            "",
            "Cierre: Gracias otra vez — y saludos a Gabriel. Te escribo con un solo siguiente paso claro.",
        ),
        alt=True,
    )

    add_pair(en, es,
        "POST-CALL — PILOT TARGET LIST",
        "DESPUÉS DE LA LLAMADA — PILOT TARGET LIST",
        "Log within 24 hours",
        "Registrar en 24 horas",
        lines(
            "Reply Notes: call date, summary, concerns, demo reaction, referral discussion.",
            "Pilot Fit: Feedback / Referral Only (unless real scenario + opt-in).",
            "Outreach Status: Call Completed.",
            "Next Follow-Up Date: +7 days if sending one-pager; +60–90 if referral may mature.",
            "Warm Intro?: update if referral path discussed.",
        ),
        lines(
            "Reply Notes: fecha, resumen, preocupaciones, reacción a la demo, conversación de referencia.",
            "Pilot Fit: Solo retroalimentación / referencia (salvo escenario real + opt-in).",
            "Outreach Status: Call Completed.",
            "Next Follow-Up Date: +7 días si envías resumen; +60–90 si la referencia puede madurar.",
            "Warm Intro?: actualizar si se discutió vía de referencia.",
        ),
    )

    add_pair(en, es,
        "TONE GUARDRAILS",
        "PAUTAS DE TONO",
        "Remember on the call",
        "Recordar en la llamada",
        lines(
            "Goal is learning, not closing a sale.",
            "Listen more than pitch; one clear next step at close.",
            "No promises on brand introductions, funding, or guaranteed matches.",
            "Demo is screen-share sample data — not production onboarding.",
            "Spanish-first; match prior email tone: personal, concise, low-pressure.",
        ),
        lines(
            "El objetivo es aprender, no cerrar una venta.",
            "Escuchar más que presentar; un solo siguiente paso al cerrar.",
            "Sin promesas de intros a marcas, financiamiento ni matches garantizados.",
            "La demo es compartir pantalla con datos de ejemplo — no onboarding a producción.",
            "Español primero; tono del email previo: personal, conciso, sin presión.",
        ),
        alt=True,
    )

    return en, es


def flatten_with_block_headers(blocks: list[tuple[str, list[dict]]]) -> list[dict]:
    """Insert a block title section before each resource group (for combined docs)."""
    out: list[dict] = []
    for i, (title, sections) in enumerate(blocks):
        out.append(
            {
                "header": title,
                "alt": i % 2 == 1,
                "subject": "",
                "subject_label": "Subject: ",
                "body": [],
                "header_only": True,
            }
        )
        out.extend(sections)
    return out


def build_combined_sections(lang: str) -> list[dict]:
    warm_en, warm_es = build_warm_intro_sections()
    reply_en, reply_es = build_reply_playbook_sections()
    accept_en, accept_es = build_acceptance_sections()
    call_en, call_es = build_pilot_call_script_sections()

    if lang == "en":
        return flatten_with_block_headers(
            [
                ("WARM INTRODUCTION TEMPLATES", warm_en),
                ("REPLY PLAYBOOK", reply_en),
                ("PILOT ACCEPTANCE CRITERIA", accept_en),
                ("PILOT CALL SCRIPT", call_en),
            ]
        )
    return flatten_with_block_headers(
        [
            ("PLANTILLAS DE INTRODUCCIÓN", warm_es),
            ("GUÍA DE RESPUESTAS", reply_es),
            ("CRITERIOS DE ACEPTACIÓN DEL PILOTO", accept_es),
            ("GUIÓN DE LLAMADA PILOTO", call_es),
        ]
    )


def append_sections_to_table(table, sections: list[dict]) -> None:
    for section in sections:
        header_row = table.add_row()
        clear_cell(header_row.cells[0])
        add_section_header(header_row.cells[0], section["header"], alt=section.get("alt", False))
        if section.get("subject"):
            add_subject_line(
                header_row.cells[0],
                section["subject"],
                label=section.get("subject_label", "Subject: "),
            )

        if section.get("header_only"):
            continue

        body_row = table.add_row()
        clear_cell(body_row.cells[0])
        for text, bold_phrases in section.get("body", []):
            if text == "":
                add_blank_line(body_row.cells[0])
            else:
                add_body_paragraph(body_row.cells[0], text, bold_phrases=bold_phrases)


def populate_table_from_sections(doc: Document, sections: list[dict], title: str) -> None:
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    table = doc.tables[0]
    remove_all_table_rows(table)

    title_row = table.add_row()
    clear_cell(title_row.cells[0])
    add_section_header(title_row.cells[0], title)

    append_sections_to_table(table, sections)


def generate_doc(sections: list[dict], title: str, output_path: Path) -> None:
    template = ensure_local_template()
    shutil.copy2(template, output_path)
    doc = Document(str(output_path))
    populate_table_from_sections(doc, sections, title)
    doc.save(str(output_path))


def write_resource_pair(
    drive: Path,
    base_name: str,
    title_en: str,
    title_es: str,
    build_sections,
) -> list[Path]:
    en_sections, es_sections = build_sections()
    en_path = drive / f"{base_name} - English.docx"
    es_path = drive / f"{base_name} - Spanish.docx"
    generate_doc(en_sections, title_en, en_path)
    generate_doc(es_sections, title_es, es_path)
    return [en_path, es_path]


def main() -> None:
    drive = find_drive_folder()
    written: list[Path] = []

    written += write_resource_pair(
        drive,
        "Dealality - GTM Pilot Resources (Jul 2026)",
        "DEALALITY GTM PILOT RESOURCES — JUL 2026",
        "RECURSOS GTM PILOTO DEALALITY — JUL 2026",
        lambda: (build_combined_sections("en"), build_combined_sections("es")),
    )

    written += write_resource_pair(
        drive,
        "Dealality - Warm Introduction Templates (Pilot Wave 1)",
        "WARM INTRODUCTION TEMPLATES — PILOT WAVE 1",
        "PLANTILLAS DE INTRODUCCIÓN — PILOTO OLA 1",
        build_warm_intro_sections,
    )
    written += write_resource_pair(
        drive,
        "Dealality - Reply Playbook (Pilot Wave 1)",
        "REPLY PLAYBOOK — PILOT WAVE 1",
        "GUÍA DE RESPUESTAS — PILOTO OLA 1",
        build_reply_playbook_sections,
    )
    written += write_resource_pair(
        drive,
        "Dealality - Pilot Acceptance Criteria (Pilot Wave 1)",
        "PILOT ACCEPTANCE CRITERIA — PILOT WAVE 1",
        "CRITERIOS DE ACEPTACIÓN DEL PILOTO — OLA 1",
        build_acceptance_sections,
    )
    written += write_resource_pair(
        drive,
        "Dealality - Pilot Call Script (Pilot Wave 1)",
        "PILOT CALL SCRIPT — PILOT WAVE 1",
        "GUIÓN DE LLAMADA PILOTO — OLA 1",
        build_pilot_call_script_sections,
    )
    written += write_resource_pair(
        drive,
        "Dealality - Daniel Shamah Call Package (Jul 9 2026)",
        "DANIEL SHAMAH CALL PACKAGE — JUL 9 2026",
        "PAQUETE DE LLAMADA DANIEL SHAMAH — 9 JUL 2026",
        build_daniel_shamah_call_package_sections,
    )

    # Remove legacy single-file bilingual versions if present
    legacy_names = [
        "Dealality - GTM Pilot Resources (Jul 2026).docx",
        "Dealality - Warm Introduction Templates (Pilot Wave 1).docx",
        "Dealality - Reply Playbook (Pilot Wave 1).docx",
        "Dealality - Pilot Acceptance Criteria (Pilot Wave 1).docx",
        "Dealality - Pilot Call Script (Pilot Wave 1).docx",
    ]
    for name in legacy_names:
        legacy = drive / name
        if legacy.exists():
            legacy.unlink()
            print(f"Removed legacy bilingual file: {legacy}")

    print("Wrote:")
    for p in written:
        print(f"  {p} ({p.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
