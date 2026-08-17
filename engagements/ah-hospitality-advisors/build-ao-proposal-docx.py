"""Build AH collaboration framework as AO Proposal Word doc from blank template."""
from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

BLANK = Path(
    r"C:\Users\joand\OneDrive\Documents\Personal Files\Consulting Projects"
    r"\AO Hospitality Advisors\Choice Hotels\Choice Hotels (CALA)"
    r"\Distribution Project"
    r"\CALA Sales ^0 Distribution Roadmap Business Proposal - Jan_29_2025 - DRAFT - BLANK.docx"
)
OUT_DIR = Path(__file__).resolve().parent
OUT_AO = Path(
    r"C:\Users\joand\OneDrive\Documents\Personal Files\Consulting Projects"
    r"\AO Hospitality Advisors\AH Hospitality Advisors"
)
FNAME = (
    "AH Hospitality Advisors - Commercial Performance Hub "
    "Collaboration Framework - Jul_16_2026 - DRAFT.docx"
)

RED = RGBColor(0xC0, 0x00, 0x00)
FONT = "Poppins"
SIZE = Pt(10)


def set_run(run, text, bold=False, red=False):
    run.text = text
    run.bold = bold
    run.font.name = FONT
    run.font.size = SIZE
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:eastAsia"), FONT)
    if red:
        run.font.color.rgb = RED
    else:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def clear_cell(cell):
    tc = cell._tc
    for child in list(tc):
        if child.tag == qn("w:p"):
            tc.remove(child)
    tc.append(OxmlElement("w:p"))


def add_para(cell, text="", bold=False, red=False, space_after=6, indent=None):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent is not None:
        p.paragraph_format.left_indent = Pt(indent)
    if text:
        run = p.add_run()
        set_run(run, text, bold=bold, red=red)
    return p


def add_heading(cell, text):
    return add_para(cell, text, bold=True, red=True, space_after=8)


def add_body(cell, text, bold=False):
    return add_para(cell, text, bold=bold, red=False, space_after=6)


def add_bullet(cell, text):
    return add_para(cell, "• " + text, space_after=3, indent=12)


def add_bold_body(cell, lead, rest):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run()
    set_run(r1, lead, bold=True)
    r2 = p.add_run()
    set_run(r2, rest)
    return p


def main():
    OUT_AO.mkdir(parents=True, exist_ok=True)
    out1 = OUT_DIR / FNAME
    out2 = OUT_AO / FNAME
    shutil.copy2(BLANK, out1)
    doc = Document(str(out1))

    # Greeting
    gcell = doc.tables[1].rows[0].cells[0]
    clear_cell(gcell)
    p = gcell.paragraphs[0]
    set_run(p.add_run(), "Dear Dean and Osama,")
    p2 = gcell.add_paragraph()
    set_run(p2.add_run(), "A&H Hospitality Advisors", bold=True)

    # Body
    bcell = doc.tables[1].rows[1].cells[0]
    clear_cell(bcell)
    first = bcell.paragraphs[0]
    set_run(
        first.add_run(),
        "Thank you for the continued conversations on building a tech-enabled "
        "commercial performance hub for A&H clients. Below is a working draft "
        "collaboration framework for discussion — not a binding agreement — "
        "reflecting our alignment calls on 7 July and 16 July 2026.",
    )

    add_heading(bcell, "Executive Summary")
    add_body(
        bcell,
        "AO Hospitality Advisors proposes to subcontract to A&H Hospitality Advisors "
        "to design, build, and operate a commercial performance hub for A&H hotel "
        "clients, starting with one incubation property (current owner-operated "
        "Courtyard mandate).",
    )
    add_body(
        bcell,
        "Goal: give owners and A&H advisors a single, trusted view of the KPIs that "
        "matter for commercial decisions — without A&H becoming a software company.",
    )
    add_bold_body(
        bcell,
        "Dashboard #1 (agreed direction, 16 July): ",
        "Actuals · Budget · Forecast by segment — with segmentation + channel as the "
        "foundation. Pace is intentionally not the first dashboard: a single pace "
        "snapshot is weak until history accumulates; ABF by segment is more static "
        "and usable sooner.",
    )
    add_body(
        bcell,
        "Dealality (AO’s deal-flow platform) remains out of scope for this engagement. "
        "The commercial hub must stand alone — clear costs, portable data/config — "
        "under the A&H brand unless A&H decides otherwise.",
    )

    add_heading(bcell, "Engagement Objectives")
    add_bullet(
        bcell,
        "Establish a scalable data factory (ingest → clean → standardize → dashboard) "
        "using standardized upload first.",
    )
    add_bullet(
        bcell,
        "Deliver Dashboard #1: Actuals · Budget · Forecast by segment "
        "(with channel / direct vs indirect where available).",
    )
    add_bullet(
        bcell,
        "Define ownership, hosting, EU compliance, and commercial terms so both "
        "parties can scale to additional properties without rework.",
    )
    add_bullet(
        bcell,
        "Keep Phase 1 narrow; add pace, overlays, and Stairway-class snapshot "
        "history only after the foundation is live.",
    )

    add_heading(bcell, "Proposed Commercial Structure")
    add_body(
        bcell,
        "A&H = prime contractor to the hotel owner. AO = subcontractor to A&H, "
        "working under the A&H flag.",
    )
    add_bullet(
        bcell,
        "Contracting: A&H ↔ Client (MSA/SOW); A&H ↔ AO (subcontractor agreement + "
        "SOW per phase).",
    )
    add_bullet(
        bcell,
        "Branding (near-term default): A&H Hospitality Advisors on the commercial / "
        "performance hub. No Dealality branding on this deliverable.",
    )
    add_bullet(
        bcell,
        "Invoicing: A&H invoices the client; A&H pays AO per subcontract milestones.",
    )
    add_bullet(bcell, "Entity: AO = A-O entity (US); A&H = UK entity.")

    add_heading(bcell, "Ownership Summary")
    add_body(
        bcell,
        "There is no single owner of “the dashboard.” Ownership splits across three "
        "layers — by design:",
    )
    add_bullet(bcell, "A&H owns the client — relationship, advisory, invoicing.")
    add_bullet(
        bcell,
        "The hotel owner owns their data — actuals, forecasts, PMS exports, etc.",
    )
    add_bullet(
        bcell,
        "AO owns the platform — data factory, parsers, reusable UI shell, and code. "
        "A&H receives a non-exclusive license to deploy for A&H clients during the "
        "partnership.",
    )
    add_bullet(
        bcell,
        "The configured dashboard instance is a licensed deliverable for the A&H "
        "mandate — not a software product the owner buys outright.",
    )
    add_body(
        bcell,
        "In one line: A&H owns the client · the owner owns the data · AO owns the "
        "technology · the client licenses the running instance.",
    )

    add_heading(bcell, "If A&H and AO Part Ways")
    add_bullet(
        bcell,
        "Existing dashboards for active A&H mandates may keep running under a "
        "maintenance arrangement (AO retainer or agreed handover).",
    )
    add_bullet(
        bcell,
        "AO platform source code is not handed over by default; A&H may keep a "
        "maintenance license/retainer or freeze the deployment as-is.",
    )
    add_bullet(
        bcell,
        "Platform buyout is a separate commercial conversation, not included in "
        "build fees.",
    )
    add_bullet(
        bcell,
        "On exit, A&H receives a handover package: runbook, source map, and export "
        "of the client’s own data and configuration values (not AO source code).",
    )

    add_heading(bcell, "Guiding Principles")
    add_bullet(bcell, "One dashboard, phased — under-promise, over-deliver.")
    add_bullet(
        bcell,
        "Foundation before pace — segmentation (+ channel) → ABF by segment → then "
        "pace / overlays.",
    )
    add_bullet(
        bcell,
        "Data factory first — standardized upload short-term; live PMS integrations "
        "later.",
    )
    add_bullet(
        bcell,
        "Client-defined rules — segmentation, fiscal calendar, mapping library with "
        "an Other catchall.",
    )
    add_bullet(
        bcell,
        "Stand-alone platform — not mixed into Dealality; portable enough for another "
        "operator later.",
    )
    add_bullet(
        bcell,
        "EU-first compliance — design for EU GDPR; EU AI Act only when AI features "
        "are introduced (MVP is rules/code-based).",
    )

    add_heading(bcell, "Scope of Work")

    add_body(bcell, "Phase 0 — Discovery & Signed Wireframe (fixed fee)", bold=True)
    add_body(
        bcell,
        "Objective: Lock Dashboard #1 = Actuals · Budget · Forecast by segment; "
        "complete source map, segmentation/channel rules, and success criteria "
        "before build.",
    )
    add_bullet(
        bcell,
        "Data & source inventory including 2–3 real sample files (A&H leads; AO "
        "documents).",
    )
    add_bullet(
        bcell,
        "Segmentation dictionary + channel dimension (direct / indirect / other) + "
        "YoY mapping rules.",
    )
    add_bullet(
        bcell,
        "Code/keyword → segment mapping library with Other catchall.",
    )
    add_bullet(
        bcell,
        "Fiscal calendar & budget/forecast input rules (entered once, used "
        "everywhere).",
    )
    add_bullet(bcell, "Clickable wireframe / mockup sign-off for ABF-by-segment.")
    add_bullet(bcell, "Hosting region + estimated annual infra cost (EU-first).")
    add_bullet(bcell, "Phase 1 scope, timeline, and single fixed price.")
    add_body(
        bcell,
        "Exit criteria: written sign-off on wireframe + source map + mapping library "
        "outline + Phase 1 SOW.",
    )
    add_body(bcell, "Indicative planning range (not a quote): USD $3,000–$8,000.")

    add_body(
        bcell,
        "Phase 1 — MVP: Data Factory + Dashboard #1 (fixed fee — quoted after Phase 0)",
        bold=True,
    )
    add_body(
        bcell,
        "Objective: one working Actuals · Budget · Forecast by segment dashboard on "
        "the incubation hotel, with channel visibility where data allows, fed by "
        "standardized upload, hosted EU-compliant with secure A&H-branded logins.",
    )
    add_bullet(
        bcell,
        "Data factory for agreed sources only: validation, segment + channel mapping, "
        "refresh log, missed-upload alerts.",
    )
    add_bullet(
        bcell,
        "Core inputs: budget once; forecast once (transient vs group as required); "
        "segmentation config.",
    )
    add_bullet(
        bcell,
        "Commentary with audit trail; rules-based owner narrative (LLM optional in "
        "Phase 2+).",
    )
    add_bullet(bcell, "Handover runbook: who uploads what, when; how to refresh.")
    add_body(
        bcell,
        "Out of scope for MVP unless added by change order: full booking pace "
        "workspace; Stairway-class multi-year daily snapshot history; live PMS/RMS "
        "APIs; displacement calculator; portfolio roll-up; AI/LLM; additional "
        "dashboards; replacing Stairway as RMS.",
    )
    add_body(
        bcell,
        "Acceptance criteria (all required): real-file ingest; segment + channel "
        "reconciliation; wireframe match; budget/forecast inputs; commentary; "
        "narrative; logins + A&H branding; runbook walkthrough.",
    )
    add_body(
        bcell,
        "Indicative planning range (replaced by fixed quote after Phase 0): "
        "USD $15,000–$40,000.",
    )

    add_body(
        bcell,
        "Phase 2 — Enhancements & Overlays (quoted per item / retainer)",
        bold=True,
    )
    add_body(
        bcell,
        "After the ABF foundation is live: booking pace; forecast accuracy; pace + "
        "forecast + comp overlays; market/comp views; LLM narrative (with EU AI Act "
        "review); portfolio; alerting; closer-to-live integrations.",
    )
    add_body(
        bcell,
        "Indicative retainer: USD $2,000–$6,000/month AO subcontract (includes refresh "
        "monitoring, hosting, up to 4 hours/month light tweaks, monthly check-in). "
        "New sources, new dashboards, methodology changes, and APIs are change orders.",
    )

    add_body(bcell, "Phase 3 — Additional Properties", bold=True)
    add_body(
        bcell,
        "Per-property onboarding from a rate card after Phase 1 (typically lower than "
        "MVP because the factory exists). Ready-to-push checklist: source map, "
        "segmentation config, A&H training, UAT, support model.",
    )

    add_heading(bcell, "Longer Horizon — Stairway-Class Capability (Not Phase 1)")
    add_body(
        bcell,
        "Context (16 July): incubation owner currently uses Stairway (~€5,000/month). "
        "Contract renews September 2026 for another year; A&H has told the owner they "
        "are not ready to replace Stairway now. Realistic replace window: ~September "
        "2027, preferably with more than one client sharing platform cost. Replacing "
        "Stairway also implies A&H operational RMS capacity (e.g. revenue manager) — "
        "advisory/ops, not only software.",
    )
    add_body(
        bcell,
        "MVP does not commit to Stairway replacement. Phase 1 builds the scalable kit "
        "of parts so history/pace layers can be added without throwing away the "
        "foundation.",
    )

    add_heading(bcell, "Data Architecture (Direction)")
    add_body(
        bcell,
        "Short-term: upload / agreed extracts → data factory (validate, segment map "
        "library, channel map) → standard DB (actuals, budget, forecast, segment, "
        "channel; pace later) → dashboard layer (ABF by segment + commentary).",
    )
    add_bullet(
        bcell,
        "Segment = selling/marketing pattern (client-defined; brands/RMS vendors "
        "often disagree).",
    )
    add_bullet(
        bcell,
        "Channel = how the booking arrived (direct vs indirect); owner-critical for "
        "“value of the deal.”",
    )
    add_bullet(
        bcell,
        "Both dimensions must reconcile to the same totals; include an Other catchall.",
    )

    add_heading(bcell, "Pricing Model Recommendation")
    add_body(
        bcell,
        "Prefer a clean first step: fixed fee + fixed timeline + signed deliverable "
        "for Phase 0/1. Revisit percentage alignment after the first deliverable works.",
    )
    add_bullet(bcell, "Phase 0: always fixed.")
    add_bullet(bcell, "Phase 1: fixed quote after Phase 0 (one number, not a range).")
    add_bullet(bcell, "Phase 2: fixed monthly retainer.")
    add_bullet(bcell, "Phase 3: fixed per-property rate card.")
    add_body(
        bcell,
        "Optional later alternative: 25–35% of A&H’s client implementation fee with "
        "floor and ceiling — only if A&H shares client pricing transparently.",
    )
    add_body(
        bcell,
        "What A&H charges the hotel owner is entirely A&H’s decision. Ongoing "
        "hosting/database for a single-property MVP is typically low hundreds to low "
        "thousands USD per year (itemize in Phase 0).",
    )

    add_heading(bcell, "Proposed Payment Structure (AO ↔ A&H)")
    add_bullet(bcell, "Currency: USD invoiced by AO to A&H.")
    add_bullet(
        bcell,
        "Phase 0: 50% on SOW signature · 50% on delivery of signed wireframe + "
        "source map + Phase 1 quote.",
    )
    add_bullet(
        bcell,
        "Phase 1: 30% on SOW signature · 40% on data factory UAT · 30% on dashboard "
        "acceptance.",
    )
    add_bullet(
        bcell,
        "Retainer: monthly in advance; 30-day notice to pause or cancel.",
    )
    add_bullet(
        bcell,
        "Payment terms: Net 15 (negotiable to Net 30). Work pauses after 15 days "
        "overdue.",
    )
    add_bullet(bcell, "No travel/expenses unless pre-approved in writing.")

    add_heading(bcell, "Operating Models After MVP")
    add_bullet(
        bcell,
        "A — Handoff: build complete; hotel/A&H uploads; AO on call for break/fix "
        "or enhancements.",
    )
    add_bullet(
        bcell,
        "B — Managed refresh: A&H or AO processes weekly feeds into the factory.",
    )
    add_bullet(
        bcell,
        "C — AO maintain & manage: ongoing ops under retainer; enhancements quoted "
        "separately.",
    )
    add_body(bcell, "First engagement should prove Model A or B before assuming C.")

    add_heading(bcell, "Hosting, Branding & Stand-Alone Stack")
    add_bullet(
        bcell,
        "Storage/platform has a clear owner (default: AO-managed infra under A&H "
        "branding).",
    )
    add_bullet(
        bcell,
        "Stand-alone environment — not shared cost pools with Dealality.",
    )
    add_bullet(
        bcell,
        "Login URL, page title, and notifications = A&H Hospitality Advisors.",
    )
    add_bullet(bcell, "Design for portability: export of client data + config.")

    add_heading(bcell, "Security, Confidentiality & EU-First Data Protection")
    add_body(
        bcell,
        "A&H’s near-term commercial clients include EU properties — design to the "
        "stricter EU bar.",
    )
    add_bullet(bcell, "AO is processor to A&H; EU hosting default for EU client data.")
    add_bullet(
        bcell,
        "Subcontract includes processor/DPA-style terms; subprocessors listed and "
        "approved.",
    )
    add_bullet(
        bcell,
        "MVP is rules/code-based; EU AI Act review before any LLM/AI feature.",
    )
    add_bullet(
        bcell,
        "AO cleans and presents data but is not the system of record; A&H/owner "
        "remain responsible for business decisions.",
    )
    add_bullet(
        bcell,
        "AO liability cap: fees paid under the applicable SOW (confirm with counsel).",
    )

    add_heading(bcell, "Governance")
    add_bullet(bcell, "Biweekly steering during build; monthly in retainer.")
    add_bullet(
        bcell,
        "Written change requests with time/cost impact; A&H gatekeeps client scope.",
    )
    add_bullet(
        bcell,
        "Points of contact: A&H commercial (Dean / Osama); AO technical (Joan).",
    )

    add_heading(bcell, "Project Team")
    add_bold_body(
        bcell,
        "Joan Dejarden — Principal / Technical Lead (AO): ",
        "Hospitality strategy and technology delivery; responsible for architecture, "
        "data factory, hosting, UI, and runbooks under the A&H flag.",
    )
    add_bold_body(
        bcell,
        "A&H Hospitality Advisors (Dean Auburn & Osama): ",
        "Commercial subject-matter experts, client relationship, KPI/segmentation "
        "definition, and owner-facing advisory.",
    )

    add_heading(bcell, "Delivery Risks & Assumptions")
    add_body(
        bcell,
        "Shared openly so A&H and AO plan Phase 0/1 with eyes open. These are "
        "management risks, not reasons to pause — they define what must be true "
        "for the fixed-fee path to work.",
    )
    add_body(bcell, "Assumptions (Phase 1 depends on these)", bold=True)
    add_bullet(
        bcell,
        "Phase 1 starts only after Phase 0 delivers 2–3 real sample export files "
        "and a signed source map.",
    )
    add_bullet(
        bcell,
        "Dashboard #1 remains Actuals · Budget · Forecast by segment (+ channel "
        "only where data is confirmed available).",
    )
    add_bullet(
        bcell,
        "Short-term ingest is standardized upload / agreed extracts, not live "
        "PMS/RMS APIs.",
    )
    add_bullet(
        bcell,
        "A&H owns the segmentation + channel dictionary; AO implements the signed "
        "mapping library (including Other).",
    )
    add_bullet(
        bcell,
        "The hub is a stand-alone stack (not mixed into Dealality), with EU hosting "
        "where required.",
    )
    add_bullet(
        bcell,
        "MVP is rules/code-based (no LLM/AI until a separate EU AI Act review).",
    )
    add_bullet(
        bcell,
        "Replacing Stairway / daily multi-year snapshot history is out of Phase 1 "
        "(~Sep 2027 horizon).",
    )
    add_bullet(
        bcell,
        "A&H remains client-facing; AO does not provide on-property revenue-management "
        "staffing.",
    )

    add_body(bcell, "Delivery risks & mitigations", bold=True)
    add_bold_body(
        bcell,
        "Data quality / messiness unknown: ",
        "Until real files are seen, effort and the Phase 1 fixed price are estimates. "
        "Mitigation: Phase 0 includes sample files; Phase 1 is one fixed quote only "
        "after file review. Late or unusable files → timeline flex or Phase 0 kill.",
    )
    add_bold_body(
        bcell,
        "Scope creep into pace / Stairway-class: ",
        "Pace and daily snapshot history are a different product class. Mitigation: "
        "written SOW boundaries; pace/overlays = Phase 2+; Stairway replacement = "
        "separate later SOW.",
    )
    add_bold_body(
        bcell,
        "Segmentation / channel disagreement: ",
        "Brand, Stairway, and owner definitions often conflict. Mitigation: A&H signs "
        "the dictionary; AO implements only the signed map; reconciliation + Other "
        "catchall required.",
    )
    add_bold_body(
        bcell,
        "Direct/indirect not available at hotel level: ",
        "Channel views are owner-critical but access is unconfirmed. Mitigation: "
        "confirm in Phase 0; if unavailable, ship ABF-by-segment without channel "
        "(or limited proxy) and quote channel as a change order.",
    )
    add_bold_body(
        bcell,
        "Client / A&H data latency: ",
        "Waiting on extracts looks like AO delay. Mitigation: day-for-day timeline "
        "extension; pause rights after extended delay.",
    )
    add_bold_body(
        bcell,
        "EU hosting / GDPR / future AI: ",
        "Wrong region or premature AI creates compliance risk. Mitigation: EU host "
        "in Phase 0; DPA-style terms; AI only after joint EU AI Act review.",
    )
    add_bold_body(
        bcell,
        "Operating-model mismatch: ",
        "If A&H expects AO to run Stairway-like RMS ops, that is staffing + advisory — "
        "not Phase 1 software. Mitigation: Phase 1 = build + handoff or light refresh "
        "(Models A/B); Model C only under explicit retainer; RMS ops remains A&H.",
    )
    add_bold_body(
        bcell,
        "Planning ranges read as quotes: ",
        "USD $15k–$40k can be treated as a firm offer before discovery. Mitigation: "
        "all Phase 1 ranges are planning anchors only; the binding number is the "
        "post–Phase 0 fixed quote.",
    )
    add_bold_body(
        bcell,
        "Solo delivery / calendar risk: ",
        "Lean delivery and AI-assisted coding do not remove QA, demos, or dependency "
        "waits. Mitigation: buffer calendar time; under-promise weeks; change control "
        "for new widgets/sources.",
    )
    add_bold_body(
        bcell,
        "Brand / PMS system changes: ",
        "Platform moves can change feed formats mid-engagement. Mitigation: adaptable "
        "parsers; format changes after Phase 0 = change order.",
    )

    add_body(bcell, "What “done” for Phase 1 does not mean", bold=True)
    add_bullet(bcell, "Not a Stairway replacement")
    add_bullet(bcell, "Not full booking-pace history")
    add_bullet(bcell, "Not live PMS integration")
    add_bullet(bcell, "Not AI-generated commercial advice")
    add_bullet(
        bcell,
        "Not AO owning the client relationship or owner RM decisions",
    )

    add_heading(bcell, "Illustrative Commercial Summary (Planning Anchors Only)")
    add_body(bcell, "Phase 0 Discovery + signed ABF wireframe — Fixed — USD $3k–$8k")
    add_body(
        bcell,
        "Phase 1 Factory + ABF-by-segment dashboard (1 hotel) — Fixed quote after "
        "Phase 0 — USD $15k–$40k planning range",
    )
    add_body(
        bcell,
        "Phase 2 Pace / overlays / enhancements — Fixed monthly retainer + per-SOW — "
        "USD $2k–$6k/mo",
    )
    add_body(
        bcell,
        "Phase 3 Additional properties — Fixed per property (rate card) — TBD after "
        "Phase 1",
    )
    add_body(
        bcell,
        "Later Stairway-class snapshot platform — Separate SOW — TBD; multi-client "
        "preferred",
    )
    add_body(
        bcell,
        "These anchors are not binding quotes. Phase 1 becomes a single fixed price "
        "in the Phase 0 deliverable.",
    )

    add_heading(bcell, "Suggested Next Steps")
    add_bullet(
        bcell,
        "Dean reviews this draft with Osama and mark up ownership, hosting, branding, "
        "and fixed-fee Phase 1.",
    )
    add_bullet(
        bcell,
        "Confirm Dashboard #1 = ABF by segment and any performance-test KPIs for v1.",
    )
    add_bullet(bcell, "Schedule a three-way regroup.")
    add_bullet(
        bcell,
        "Kick off Phase 0 with 2–3 real sample export files from the incubation hotel.",
    )
    add_body(
        bcell,
        "Interactive mockup (illustrative data): "
        "https://joandejarden-create.github.io/ah-commercial-performance-hub/",
    )

    add_heading(bcell, "Conclusion")
    add_body(
        bcell,
        "We believe starting with Actuals · Budget · Forecast by segment — on a "
        "stand-alone, EU-ready stack under the A&H flag — is the cleanest path to a "
        "valuable first deliverable, while preserving the option to add pace and "
        "Stairway-class capability over the next year.",
    )
    add_body(
        bcell,
        "This document is a working draft for discussion. I am happy to revise based "
        "on your and Osama’s comments and then convert agreed sections into a Phase 0 "
        "SOW.",
    )
    add_body(
        bcell,
        "Please don’t hesitate to contact me directly at +34 674 993 637 or via email "
        "at hello@AOHospitalityAdvisors.com should you have any questions.",
    )
    add_body(bcell, "I look forward to continuing the conversation.")
    add_para(bcell, "")
    add_body(bcell, "Best regards,")
    add_body(bcell, "Joan Dejarden")
    add_body(bcell, "AO Hospitality Advisors")

    # Document properties
    core = doc.core_properties
    core.author = "Joan Dejarden / AO Hospitality Advisors"
    core.title = (
        "AH Hospitality Advisors — Commercial Performance Hub "
        "Collaboration Framework (Draft)"
    )
    core.subject = "Working draft for discussion — Jul 16, 2026"

    doc.save(str(out1))
    shutil.copy2(out1, out2)
    print(f"Wrote: {out1}")
    print(f"Wrote: {out2}")
    print(f"Size: {out1.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
